
# utils/report_word.py — Observación con emoji 🟦 y Consideración ⚠, estilos y orden pedidos
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from datetime import datetime

def _ensure_styles(doc: Document):
    styles = doc.styles
    # Títulos para Consideración (rojo) y Observación (azul)
    if "TituloConsideracion" not in [s.name for s in styles]:
        s = styles.add_style("TituloConsideracion", WD_STYLE_TYPE.CHARACTER)
        font = s.font
        font.bold = True
        font.size = Pt(10)
        font.color.rgb = RGBColor(200, 0, 0)
        font.name = "Segoe UI"
    if "TituloObservacion" not in [s.name for s in styles]:
        s = styles.add_style("TituloObservacion", WD_STYLE_TYPE.CHARACTER)
        font = s.font
        font.bold = True
        font.size = Pt(10)
        font.color.rgb = RGBColor(0, 90, 200)
        font.name = "Segoe UI"
    if "NormalBody" not in [s.name for s in styles]:
        s = styles.add_style("NormalBody", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Segoe UI"
        s.font.size = Pt(10)

def build_word(title: str, steps: list, out_path: str):
    doc = Document()
    _ensure_styles(doc)

    # Portada mínima
    p = doc.add_paragraph()
    run = p.add_run(title or "Incidencia")
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph(datetime.now().strftime("Generado: %Y-%m-%d %H:%M:%S"))
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Pasos enumerados
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    for idx, st in enumerate(steps, start=1):
        paso = doc.add_paragraph()
        paso_run = paso.add_run(f"Paso {idx}")
        paso_run.bold = True
        paso_run.font.size = Pt(12)

        desc = (st.get("desc") or "").strip()
        cons = (st.get("consideraciones") or "").strip()
        obs  = (st.get("observacion") or "").strip()
        shots = [s for s in (st.get("shots") or []) if s]

        if desc:
            doc.add_paragraph(desc, style="NormalBody")

        for shot in shots:
            try:
                doc.add_picture(str(shot), width=Inches(6.5))
            except Exception:
                doc.add_paragraph(f"[No se pudo insertar imagen: {shot}]")

        if cons:
            p = doc.add_paragraph()
            p.add_run("⚠️ Consideración: ").style = "TituloConsideracion"
            p.add_run(cons)

        if obs:
            p = doc.add_paragraph()
            p.add_run("✅ Observación: ").style = "TituloObservacion"
            p.add_run(obs)

        doc.add_paragraph("")

    doc.save(out_path)
