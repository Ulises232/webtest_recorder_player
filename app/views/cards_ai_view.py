"""Build the UI that orchestrates card selection and AI generations."""

from __future__ import annotations

import json
import re
import threading
from datetime import datetime
from html import escape
from importlib import resources
from typing import Callable, Dict, List, Optional
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import ttkbootstrap as tb
from ttkbootstrap.constants import *  # noqa: F401,F403
from ttkbootstrap.widgets import DateEntry

from app.controllers.card_ai_controller import CardAIController
from app.dtos.card_ai_dto import CardAIGenerationResultDTO, CardAIHistoryEntryDTO, CardDTO


TYPE_CHOICES = ("INCIDENCIA", "MEJORA", "HU")

HTML_TEMPLATE_PACKAGE = "app.templates"
HTML_TEMPLATE_NAME = "card_generation.html"


def _format_datetime(value: Optional[datetime]) -> str:
    """Return a friendly formatted datetime string."""

    if not value:
        return ""
    return value.strftime("%Y-%m-%d %H:%M")


def _calculate_completeness(fields: Dict[str, str]) -> int:
    """Compute the completeness percentage mirroring the service logic."""

    keys = ["descripcion", "analisis", "recomendaciones", "cosas_prevenir", "info_adicional"]
    values = [fields.get(key, "").strip() for key in keys]
    filled = sum(1 for value in values if value)
    return round(100 * filled / len(values)) if values else 0


def _progress_style(progress: tb.Progressbar, percentage: int) -> None:
    """Adjust the progress bar style according to the completion percentage."""

    if percentage < 34:
        progress.configure(bootstyle="danger")
    elif percentage < 67:
        progress.configure(bootstyle="warning")
    else:
        progress.configure(bootstyle="success")


def _export_json(content: Dict[str, object], parent: tk.Misc) -> None:
    """Prompt for a JSON path and save the provided document."""

    path = filedialog.asksaveasfilename(
        parent=parent,
        defaultextension=".json",
        filetypes=[("JSON", "*.json")],
        title="Exportar resultado a JSON",
    )
    if not path:
        return
    try:
        with open(path, "w", encoding="utf-8") as handle:
            json.dump(content, handle, ensure_ascii=False, indent=2)
    except OSError as exc:
        messagebox.showerror("Error", f"No se pudo exportar el JSON:\n{exc}")
    else:
        messagebox.showinfo("Exportado", f"Archivo guardado en:\n{path}")


def _export_markdown(content: Dict[str, object], parent: tk.Misc) -> None:
    """Transform the JSON into Markdown and persist it to disk."""

    path = filedialog.asksaveasfilename(
        parent=parent,
        defaultextension=".md",
        filetypes=[("Markdown", "*.md")],
        title="Exportar resultado a Markdown",
    )
    if not path:
        return
    try:
        lines: List[str] = ["# Documento generado"]
        for key, value in content.items():
            lines.append(f"\n## {key}")
            if isinstance(value, list):
                for item in value:
                    lines.append(f"- {item}")
            else:
                lines.append(str(value))
        with open(path, "w", encoding="utf-8") as handle:
            handle.write("\n".join(lines))
    except OSError as exc:
        messagebox.showerror("Error", f"No se pudo exportar el Markdown:\n{exc}")
    else:
        messagebox.showinfo("Exportado", f"Archivo guardado en:\n{path}")


def _export_docx(content: Dict[str, object], parent: tk.Misc) -> None:
    """Generate a minimal DOCX using python-docx if available."""

    try:
        from docx import Document  # type: ignore
    except ImportError:
        messagebox.showerror(
            "Dependencia faltante",
            "No se encontró la librería 'python-docx'. Instala la dependencia para exportar a DOCX.",
        )
        return

    path = filedialog.asksaveasfilename(
        parent=parent,
        defaultextension=".docx",
        filetypes=[("Documento Word", "*.docx")],
        title="Exportar resultado a DOCX",
    )
    if not path:
        return

    try:
        document = Document()
        document.add_heading("Documento DDE/HU", level=1)
        for key, value in content.items():
            document.add_heading(str(key), level=2)
            if isinstance(value, list):
                for item in value:
                    document.add_paragraph(str(item), style="List Bullet")
            else:
                document.add_paragraph(str(value))
        document.save(path)
    except Exception as exc:  # pragma: no cover - depende de la librería externa
        messagebox.showerror("Error", f"No se pudo exportar el DOCX:\n{exc}")
    else:
        messagebox.showinfo("Exportado", f"Archivo guardado en:\n{path}")


def _load_html_template() -> str:
    """Retrieve the HTML template used for document exports."""

    try:
        template_resource = resources.files(HTML_TEMPLATE_PACKAGE).joinpath(HTML_TEMPLATE_NAME)
    except ModuleNotFoundError as exc:  # pragma: no cover - configuración inválida
        raise FileNotFoundError("No se encontró el paquete de plantillas HTML.") from exc

    try:
        with template_resource.open("r", encoding="utf-8") as handle:
            return handle.read()
    except FileNotFoundError as exc:
        raise FileNotFoundError("No se encontró la plantilla HTML configurada.") from exc
    except OSError as exc:  # pragma: no cover - errores del sistema de archivos
        raise OSError("Ocurrió un error al leer la plantilla HTML.") from exc


def _format_html_value(value: object) -> str:
    """Convert plain text or lists into HTML-friendly blocks."""

    if isinstance(value, list):
        items = [escape(str(item)) for item in value if str(item).strip()]
        if not items:
            return "&nbsp;"
        return "<ul>" + "".join(f"<li>{item}</li>" for item in items) + "</ul>"

    if value is None:
        return "&nbsp;"

    text = escape(str(value).strip())
    if not text:
        return "&nbsp;"
    return text.replace("\n", "<br />")


def _build_html_document(content: Dict[str, object]) -> str:
    """Render the HTML export by injecting the AI response into the template."""

    template = _load_html_template()
    placeholders: Dict[str, str] = {
        "titulo": _format_html_value(content.get("titulo")),
        "tipo": _format_html_value(content.get("tipo")),
        "fecha": _format_html_value(content.get("fecha")),
        "hora_inicio": _format_html_value(content.get("hora_inicio")),
        "hora_fin": _format_html_value(content.get("hora_fin")),
        "lugar": _format_html_value(content.get("lugar")),
        "descripcion": _format_html_value(content.get("descripcion")),
        "que_necesitas": _format_html_value(content.get("que_necesitas")),
        "para_que_lo_necesitas": _format_html_value(content.get("para_que_lo_necesitas")),
        "como_lo_necesitas": _format_html_value(content.get("como_lo_necesitas")),
        "requerimientos_funcionales": _format_html_value(content.get("requerimientos_funcionales")),
        "requerimientos_especiales": _format_html_value(content.get("requerimientos_especiales")),
        "criterios_aceptacion": _format_html_value(content.get("criterios_aceptacion")),
    }

    rendered = template
    for key, value in placeholders.items():
        rendered = rendered.replace(f"{{{{{key}}}}}", value)

    rendered = re.sub(r"\{\{[a-zA-Z0-9_]+\}\}", "&nbsp;", rendered)
    return rendered


def _export_html(content: Dict[str, object], parent: tk.Misc) -> None:
    """Generate an HTML file using the configurable template."""

    path = filedialog.asksaveasfilename(
        parent=parent,
        defaultextension=".html",
        filetypes=[("HTML", "*.html")],
        title="Exportar resultado a HTML",
    )
    if not path:
        return

    try:
        document = _build_html_document(content)
        with open(path, "w", encoding="utf-8") as handle:
            handle.write(document)
    except FileNotFoundError as exc:
        messagebox.showerror("Error", str(exc))
    except OSError as exc:
        messagebox.showerror("Error", f"No se pudo exportar el HTML:\n{exc}")
    else:
        messagebox.showinfo("Exportado", f"Archivo guardado en:\n{path}")


def _show_history(parent: tk.Misc, controller: CardAIController, card_id: int) -> None:
    """Display a modal window with the generation history for a card."""

    try:
        history_entries = controller.list_history(card_id)
    except RuntimeError as exc:
        messagebox.showerror("Error", str(exc))
        return

    win = tb.Toplevel(parent)
    win.title("Historial de generación")
    win.geometry("860x560")
    win.transient(parent)
    win.grab_set()

    frame = tb.Frame(win, padding=12)
    frame.pack(fill=BOTH, expand=YES)

    columns = ("fecha", "modelo", "completitud", "mejor", "dde")
    tree = ttk.Treeview(frame, columns=columns, show="headings", height=12)
    tree.heading("fecha", text="Fecha")
    tree.heading("modelo", text="Modelo")
    tree.heading("completitud", text="Completitud")
    tree.heading("mejor", text="Mejor respuesta")
    tree.heading("dde", text="DDE generada")
    tree.column("fecha", width=180)
    tree.column("modelo", width=200)
    tree.column("completitud", width=120, anchor="center")
    tree.column("mejor", width=140, anchor="center")
    tree.column("dde", width=140, anchor="center")
    tree.tag_configure("best", background="#e6f4ea")
    tree.tag_configure("dde", background="#fff7e6")

    scrollbar = ttk.Scrollbar(frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=scrollbar.set)
    tree.pack(side=LEFT, fill=BOTH, expand=YES)
    scrollbar.pack(side=RIGHT, fill=Y)

    detail = tk.Text(win, height=10, wrap="word")
    detail.pack(fill=BOTH, expand=YES, padx=12, pady=(0, 12))

    entries_map: Dict[str, CardAIHistoryEntryDTO] = {}
    managed_buttons: List[tk.Widget] = []
    best_toggle_button: Optional[tk.Button] = None

    def get_selected_key() -> Optional[str]:
        """Return the identifier for the currently selected row."""

        selection = tree.selection()
        if not selection:
            return None
        return selection[0]

    def get_selected_entry() -> Optional[CardAIHistoryEntryDTO]:
        """Return the DTO attached to the current selection."""

        key = get_selected_key()
        if not key:
            return None
        return entries_map.get(key)

    def set_buttons_state() -> None:
        """Enable or disable the history actions according to the selection."""

        entry = get_selected_entry()
        state = tk.NORMAL if entry else tk.DISABLED
        for button in managed_buttons:
            if button.winfo_exists():
                button.configure(state=state)
        if not best_toggle_button or not best_toggle_button.winfo_exists():
            return
        if not entry:
            best_toggle_button.configure(
                state=tk.DISABLED,
                text="Marcar como mejor",
                bootstyle=SUCCESS,
            )
            return
        if entry.output.isBest:
            best_toggle_button.configure(
                state=tk.NORMAL,
                text="Quitar marca mejor",
                bootstyle=WARNING,
            )
        else:
            best_toggle_button.configure(
                state=tk.NORMAL,
                text="Marcar como mejor",
                bootstyle=SUCCESS,
            )

    def on_select(event: tk.Event | None) -> None:
        """Render the JSON content for the selected history entry."""

        entry = get_selected_entry()
        detail.configure(state="normal")
        detail.delete("1.0", "end")
        if not entry:
            detail.configure(state="disabled")
            if mark_dde_button.winfo_exists():
                mark_dde_button.configure(text="Marcar DDE generada", bootstyle=INFO)
            set_buttons_state()
            return
        pretty = json.dumps(entry.output.content, ensure_ascii=False, indent=2)
        detail.insert("1.0", pretty)
        detail.configure(state="disabled")
        if entry.output.ddeGenerated:
            mark_dde_button.configure(text="Quitar marca DDE", bootstyle=WARNING)
        else:
            mark_dde_button.configure(text="Marcar DDE generada", bootstyle=INFO)
        set_buttons_state()

    def export_selected_json() -> None:
        """Export the selected history entry using the JSON helper."""

        entry = get_selected_entry()
        if not entry:
            return
        _export_json(entry.output.content, win)

    def export_selected_markdown() -> None:
        """Export the selected history entry to Markdown format."""

        entry = get_selected_entry()
        if not entry:
            return
        _export_markdown(entry.output.content, win)

    def export_selected_docx() -> None:
        """Export the selected history entry to DOCX format."""

        entry = get_selected_entry()
        if not entry:
            return
        _export_docx(entry.output.content, win)

    def export_selected_html() -> None:
        """Export the selected history entry to HTML format."""

        entry = get_selected_entry()
        if not entry:
            return
        _export_html(entry.output.content, win)

    def delete_selected_entry() -> None:
        """Delete the selected history entry after confirmation."""

        nonlocal history_entries
        key = get_selected_key()
        entry = get_selected_entry()
        if not key or not entry:
            return
        if not messagebox.askyesno(
            "Eliminar resultado",
            "¿Deseas eliminar el resultado seleccionado del historial?",
        ):
            return
        try:
            controller.delete_output(entry.output.outputId)
        except RuntimeError as exc:
            messagebox.showerror("Error", str(exc))
            return
        history_entries = [
            item for item in history_entries if item.output.outputId != entry.output.outputId
        ]
        populate_tree(history_entries)
        messagebox.showinfo("Historial", "El resultado se eliminó correctamente.")

    def toggle_selected_best_flag() -> None:
        """Toggle the preferred flag for the selected entry."""

        nonlocal history_entries
        entry = get_selected_entry()
        if not entry:
            return
        target_state = not entry.output.isBest
        try:
            if target_state:
                controller.mark_output_as_best(entry.output.outputId)
            else:
                controller.clear_output_best_flag(entry.output.outputId)
        except RuntimeError as exc:
            messagebox.showerror("Error", str(exc))
            return
        try:
            history_entries = controller.list_history(card_id)
        except RuntimeError as exc:
            messagebox.showerror("Error", str(exc))
            return
        populate_tree(history_entries, selected_output=entry.output.outputId)
        if target_state:
            messagebox.showinfo("Historial", "Se marcó la respuesta como mejor opción.")
        else:
            messagebox.showinfo("Historial", "Se quitó la marca de mejor respuesta.")

    def toggle_selected_dde() -> None:
        """Toggle the DDE generated flag for the selected entry."""

        nonlocal history_entries
        entry = get_selected_entry()
        if not entry:
            return
        target_state = not entry.output.ddeGenerated
        try:
            controller.mark_output_dde_generated(entry.output.outputId, target_state)
        except RuntimeError as exc:
            messagebox.showerror("Error", str(exc))
            return
        try:
            history_entries = controller.list_history(card_id)
        except RuntimeError as exc:
            messagebox.showerror("Error", str(exc))
            return
        populate_tree(history_entries, selected_output=entry.output.outputId)
        if target_state:
            messagebox.showinfo("Historial", "Se marcó la salida como DDE generada.")
        else:
            messagebox.showinfo("Historial", "Se quitó la marca de DDE generada.")

    actions = tb.Frame(win, padding=(12, 0, 12, 12))
    actions.pack(fill=X)

    export_json_button = tb.Button(
        actions,
        text="Exportar JSON",
        bootstyle=SECONDARY,
        command=export_selected_json,
        state=tk.DISABLED,
    )
    export_json_button.pack(side=LEFT)
    managed_buttons.append(export_json_button)

    export_md_button = tb.Button(
        actions,
        text="Exportar Markdown",
        bootstyle=SECONDARY,
        command=export_selected_markdown,
        state=tk.DISABLED,
    )
    export_md_button.pack(side=LEFT, padx=6)
    managed_buttons.append(export_md_button)

    export_docx_button = tb.Button(
        actions,
        text="Exportar DOCX",
        bootstyle=SECONDARY,
        command=export_selected_docx,
        state=tk.DISABLED,
    )
    export_docx_button.pack(side=LEFT, padx=6)
    managed_buttons.append(export_docx_button)

    export_html_button = tb.Button(
        actions,
        text="Exportar HTML",
        bootstyle=SECONDARY,
        command=export_selected_html,
        state=tk.DISABLED,
    )
    export_html_button.pack(side=LEFT, padx=6)
    managed_buttons.append(export_html_button)

    best_toggle_button = tb.Button(
        actions,
        text="Marcar como mejor",
        bootstyle=SUCCESS,
        command=toggle_selected_best_flag,
        state=tk.DISABLED,
    )
    best_toggle_button.pack(side=RIGHT, padx=(0, 6))

    mark_dde_button = tb.Button(
        actions,
        text="Marcar DDE generada",
        bootstyle=INFO,
        command=toggle_selected_dde,
        state=tk.DISABLED,
    )
    mark_dde_button.pack(side=RIGHT, padx=(0, 6))
    managed_buttons.append(mark_dde_button)

    delete_button = tb.Button(
        actions,
        text="Eliminar",
        bootstyle=DANGER,
        command=delete_selected_entry,
        state=tk.DISABLED,
    )
    delete_button.pack(side=RIGHT, padx=(0, 6))
    managed_buttons.append(delete_button)

    def populate_tree(entries: List[CardAIHistoryEntryDTO], selected_output: Optional[int] = None) -> None:
        """Refresh the grid using the provided history collection."""

        entries_map.clear()
        tree.delete(*tree.get_children(""))
        selected_item: Optional[str] = None
        for entry in entries:
            completeness = entry.input.completenessPct if entry.input else 0
            item_id = str(entry.output.outputId)
            tags_list: List[str] = []
            if entry.output.isBest:
                tags_list.append("best")
            if entry.output.ddeGenerated:
                tags_list.append("dde")
            tree.insert(
                "",
                "end",
                iid=item_id,
                values=(
                    _format_datetime(entry.output.createdAt),
                    entry.output.llmModel or "",
                    f"{completeness}%",
                    "Sí" if entry.output.isBest else "No",
                    "Sí" if entry.output.ddeGenerated else "No",
                ),
                tags=tuple(tags_list),
            )
            entries_map[item_id] = entry
            if selected_output and entry.output.outputId == selected_output:
                selected_item = item_id

        if selected_item:
            tree.selection_set(selected_item)
        elif entries:
            first = tree.get_children("")
            if first:
                tree.selection_set(first[0])
        else:
            detail.configure(state="normal")
            detail.delete("1.0", "end")
            detail.configure(state="disabled")

        on_select(None)

    tree.bind("<<TreeviewSelect>>", on_select)
    populate_tree(history_entries)

    win.wait_window()


def _show_generation_result(
    parent: tk.Misc,
    controller: CardAIController,
    card: CardDTO,
    result: CardAIGenerationResultDTO,
) -> None:
    """Display the result window with export options."""

    win = tb.Toplevel(parent)
    win.title(f"Resultado para tarjeta {card.cardId}")
    win.geometry("880x620")
    win.transient(parent)
    win.grab_set()

    header = tb.Frame(win, padding=12)
    header.pack(fill=X)
    tb.Label(
        header,
        text=f"Tarjeta {card.cardId} - {card.title}",
        font=("Segoe UI", 12, "bold"),
    ).pack(anchor=W)
    tb.Label(
        header,
        text=f"Completitud: {result.completenessPct}% - Modelo: {result.output.llmModel or 'N/D'}",
        bootstyle=SECONDARY,
    ).pack(anchor=W, pady=(4, 0))

    text = tk.Text(win, wrap="word")
    text.pack(fill=BOTH, expand=YES, padx=12, pady=8)
    text.configure(state="normal")
    pretty = json.dumps(result.output.content, ensure_ascii=False, indent=2)
    text.insert("1.0", pretty)
    text.configure(state="disabled")

    actions = tb.Frame(win, padding=12)
    actions.pack(fill=X)

    tb.Button(
        actions,
        text="Exportar JSON",
        bootstyle=SECONDARY,
        command=lambda: _export_json(result.output.content, win),
    ).pack(side=LEFT)
    tb.Button(
        actions,
        text="Exportar Markdown",
        bootstyle=SECONDARY,
        command=lambda: _export_markdown(result.output.content, win),
    ).pack(side=LEFT, padx=6)
    tb.Button(
        actions,
        text="Exportar DOCX",
        bootstyle=SECONDARY,
        command=lambda: _export_docx(result.output.content, win),
    ).pack(side=LEFT, padx=6)
    tb.Button(
        actions,
        text="Exportar HTML",
        bootstyle=SECONDARY,
        command=lambda: _export_html(result.output.content, win),
    ).pack(side=LEFT, padx=6)

    tb.Button(
        actions,
        text="Historial",
        bootstyle=INFO,
        command=lambda: _show_history(parent, controller, card.cardId),
    ).pack(side=RIGHT)

    running = tk.BooleanVar(value=False)

    def _set_running(value: bool) -> None:
        """Enable or disable the action buttons while processing."""

        if not actions.winfo_exists():
            return

        running.set(value)
        state = tk.DISABLED if value else tk.NORMAL
        for button in actions.winfo_children():
            try:
                button.configure(state=state)
            except tk.TclError:
                continue

    def _background_call(func: Callable[[], None]) -> None:
        """Execute the provided callback in a worker thread."""

        def worker() -> None:
            try:
                func()
            finally:
                try:
                    win.after(0, lambda: _set_running(False))
                except tk.TclError:
                    _set_running(False)

        _set_running(True)
        threading.Thread(target=worker, daemon=True).start()

    def regenerate() -> None:
        """Trigger a new generation using the stored input."""

        if not messagebox.askyesno("Regenerar", "¿Deseas generar nuevamente con los mismos datos?"):
            return

        def _task() -> None:
            try:
                new_result = controller.regenerate(result.input.inputId)
            except RuntimeError as exc:
                error_message = str(exc)
                win.after(0, lambda message=error_message: messagebox.showerror("Error", message))
                return

            win.after(
                0,
                lambda: (
                    win.destroy(),
                    _show_generation_result(parent, controller, card, new_result),
                ),
            )

        _background_call(_task)

    tb.Button(actions, text="Regenerar", bootstyle=PRIMARY, command=regenerate).pack(side=RIGHT, padx=(0, 8))

    win.wait_window()


def _open_capture_form(
    root: tk.Misc,
    controller: CardAIController,
    card: CardDTO,
) -> None:
    """Render the modal form that captures the generation inputs."""

    win = tb.Toplevel(root)
    win.title(f"Captura para tarjeta {card.cardId}")
    win.geometry("840x620")
    win.transient(root)
    win.grab_set()

    vars_data = {
        "tipo": tk.StringVar(value=card.cardType or TYPE_CHOICES[0]),
        "descripcion": tk.StringVar(value=""),
        "analisis": tk.StringVar(value=""),
        "recomendaciones": tk.StringVar(value=""),
        "cosas_prevenir": tk.StringVar(value=""),
        "info_adicional": tk.StringVar(value=""),
    }

    provider_value_map: Dict[str, str] = {}
    provider_options: List[str] = []
    default_provider_key = "local"
    try:
        providers = controller.list_providers()
        default_provider_key = controller.get_default_provider_key()
    except RuntimeError as exc:
        messagebox.showwarning(
            "Configuración de IA",
            f"No fue posible cargar los proveedores de IA:\n{exc}",
        )
        providers = []

    for provider in providers:
        label = provider.displayName or provider.providerKey
        provider_value_map[label] = provider.providerKey
        provider_options.append(label)

    if not provider_options:
        provider_options = ["LM Studio (Local)"]
        provider_value_map = {provider_options[0]: "local"}
        default_provider_key = "local"

    default_provider_label = next(
        (
            label
            for label, key in provider_value_map.items()
            if key == default_provider_key
        ),
        provider_options[0],
    )

    provider_var = tk.StringVar(value=default_provider_label)
    vars_data["provider"] = provider_var

    def _as_payload() -> Dict[str, object]:
        """Collect the current state of the form."""

        return {
            "cardId": card.cardId,
            "tipo": vars_data["tipo"].get(),
            "descripcion": descripcion.get("1.0", "end").strip(),
            "analisis": analisis.get("1.0", "end").strip(),
            "recomendaciones": recomendaciones.get("1.0", "end").strip(),
            "cosasPrevenir": prevenir.get("1.0", "end").strip(),
            "infoAdicional": adicional.get("1.0", "end").strip(),
            "providerKey": provider_value_map.get(provider_var.get()),
        }

    def _update_progress(*_args: object) -> None:
        """Recalculate completeness when fields change."""

        payload = _as_payload()
        completeness = _calculate_completeness(
            {
                "descripcion": payload["descripcion"],
                "analisis": payload["analisis"],
                "recomendaciones": payload["recomendaciones"],
                "cosas_prevenir": payload["cosasPrevenir"],
                "info_adicional": payload["infoAdicional"],
            }
        )
        progress.configure(value=completeness)
        _progress_style(progress, completeness)
        status.set(f"Completitud estimada: {completeness}%")

    container = tb.Frame(win, padding=12)
    container.pack(fill=BOTH, expand=YES)

    tb.Label(
        container,
        text=f"Tarjeta {card.cardId}: {card.title}",
        font=("Segoe UI", 12, "bold"),
    ).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))

    current_row = 1

    tb.Label(container, text="Proveedor de IA").grid(row=current_row, column=0, sticky="w")
    provider_box = ttk.Combobox(
        container,
        values=provider_options,
        textvariable=provider_var,
        state="readonly",
    )
    provider_box.grid(row=current_row, column=1, sticky="we")
    current_row += 1

    tb.Label(container, text="Tipo").grid(row=current_row, column=0, sticky="w")
    tipo_box = ttk.Combobox(
        container, values=list(TYPE_CHOICES), textvariable=vars_data["tipo"], state="readonly"
    )
    tipo_box.grid(row=current_row, column=1, sticky="we")
    current_row += 1

    descripcion = tk.Text(container, height=5, wrap="word")
    analisis = tk.Text(container, height=5, wrap="word")
    recomendaciones = tk.Text(container, height=5, wrap="word")
    prevenir = tk.Text(container, height=5, wrap="word")
    adicional = tk.Text(container, height=4, wrap="word")

    labels = [
        ("Descripción", descripcion),
        ("Análisis", analisis),
        ("Recomendaciones", recomendaciones),
        ("Cosas a prevenir", prevenir),
        ("Información adicional", adicional),
    ]
    row = current_row
    for label, widget in labels:
        tb.Label(container, text=label).grid(row=row, column=0, sticky="nw", pady=(8, 0))
        widget.grid(row=row, column=1, sticky="nsew", pady=(8, 0))
        row += 1

    container.grid_columnconfigure(1, weight=1)
    for idx in range(current_row, row):
        container.grid_rowconfigure(idx, weight=1)

    status = tk.StringVar(value="Completa los campos para mejorar el resultado.")
    progress = tb.Progressbar(container, maximum=100, bootstyle="danger")
    progress.grid(row=row, column=0, columnspan=2, sticky="we", pady=(12, 0))
    tb.Label(container, textvariable=status, bootstyle=SECONDARY).grid(row=row + 1, column=0, columnspan=2, sticky="w")

    buttons = tb.Frame(container)
    buttons.grid(row=row + 2, column=0, columnspan=2, sticky="we", pady=(12, 0))

    running = tk.BooleanVar(value=False)

    def _set_running(value: bool) -> None:
        """Enable or disable the action buttons."""

        if not buttons.winfo_exists():
            return

        running.set(value)
        state = tk.DISABLED if value else tk.NORMAL
        for button in buttons.winfo_children():
            try:
                button.configure(state=state)
            except tk.TclError:
                continue

    def _background_call(func: Callable[[], None]) -> None:
        """Execute the given callable in a worker thread."""

        def worker() -> None:
            try:
                func()
            finally:
                try:
                    win.after(0, lambda: _set_running(False))
                except tk.TclError:
                    _set_running(False)

        _set_running(True)
        threading.Thread(target=worker, daemon=True).start()

    def _save_draft() -> None:
        """Save the current content as draft."""

        payload = _as_payload()

        def _task() -> None:
            try:
                controller.save_draft(payload)
            except RuntimeError as exc:
                error_message = str(exc)
                win.after(
                    0,
                    lambda message=error_message: messagebox.showerror("Error", message),
                )
                return
            win.after(0, lambda: messagebox.showinfo("Guardado", "Se guardó el borrador correctamente."))

        _background_call(_task)

    def _generate() -> None:
        """Call the controller to generate the document."""

        payload = _as_payload()
        completeness = _calculate_completeness(
            {
                "descripcion": payload["descripcion"],
                "analisis": payload["analisis"],
                "recomendaciones": payload["recomendaciones"],
                "cosas_prevenir": payload["cosasPrevenir"],
                "info_adicional": payload["infoAdicional"],
            }
        )
        if completeness < 34:
            if not messagebox.askyesno(
                "Completitud baja",
                "La completitud es menor al 34%. ¿Deseas continuar de todos modos?",
            ):
                return

        def _task() -> None:
            try:
                result = controller.generate_document(payload)
            except RuntimeError as exc:
                error_message = str(exc)
                win.after(
                    0,
                    lambda message=error_message: messagebox.showerror("Error", message),
                )
                return
            win.after(0, lambda: (win.destroy(), _show_generation_result(root, controller, card, result)))

        _background_call(_task)

    tb.Button(buttons, text="Guardar borrador", bootstyle=SECONDARY, command=_save_draft).pack(side=LEFT)
    tb.Button(buttons, text="Generar (IA)", bootstyle=PRIMARY, command=_generate).pack(side=LEFT, padx=6)
    tb.Button(buttons, text="Cancelar", bootstyle=DANGER, command=win.destroy).pack(side=RIGHT)

    for widget in (descripcion, analisis, recomendaciones, prevenir, adicional):
        widget.bind("<KeyRelease>", _update_progress, add="+")

    _update_progress()
    win.wait_window()


def build_cards_ai_view(
    root: tk.Misc,
    parent: tb.Frame,
    controller: CardAIController,
    bind_mousewheel: Callable[[tk.Widget, Callable[..., None]], None],
) -> None:
    """Populate the cards AI assistant view inside the provided frame."""

    for widget in parent.winfo_children():
        widget.destroy()

    filters_frame = tb.Frame(parent, padding=12)
    filters_frame.pack(fill=X)

    try:
        incident_type_options = controller.list_incidence_types()
    except RuntimeError as exc:
        messagebox.showwarning("Catálogos", f"No fue posible cargar los tipos de incidente:\n{exc}")
        incident_type_options = []

    try:
        status_options = controller.list_statuses()
    except RuntimeError as exc:
        messagebox.showwarning("Catálogos", f"No fue posible cargar los estatus:\n{exc}")
        status_options = []

    try:
        company_options = controller.list_companies()
    except RuntimeError as exc:
        messagebox.showwarning("Catálogos", f"No fue posible cargar las empresas:\n{exc}")
        company_options = []

    incident_type_map = {option.name: option.optionId for option in incident_type_options}
    incident_type_values = [""] + [option.name for option in incident_type_options]
    status_values = [""] + status_options
    company_map = {option.name: option.optionId for option in company_options}
    company_values = [""] + [option.name for option in company_options]

    tb.Label(filters_frame, text="Tipo incidente").grid(row=0, column=0, sticky="w")
    tipo_var = tk.StringVar(value="")
    tipo_box = ttk.Combobox(
        filters_frame,
        values=incident_type_values,
        textvariable=tipo_var,
        width=18,
        state="readonly",
    )
    tipo_box.grid(row=1, column=0, sticky="we", padx=(0, 10))
    if incident_type_values:
        tipo_box.current(0)

    tb.Label(filters_frame, text="Status").grid(row=0, column=1, sticky="w")
    status_var = tk.StringVar(value="")
    status_box = ttk.Combobox(
        filters_frame,
        values=status_values,
        textvariable=status_var,
        width=16,
        state="readonly",
    )
    status_box.grid(row=1, column=1, sticky="we", padx=(0, 10))
    if status_values:
        status_box.current(0)

    tb.Label(filters_frame, text="Empresa").grid(row=0, column=2, sticky="w")
    company_var = tk.StringVar(value="")
    company_box = ttk.Combobox(
        filters_frame,
        values=company_values,
        textvariable=company_var,
        width=22,
        state="readonly",
    )
    company_box.grid(row=1, column=2, sticky="we", padx=(0, 10))
    if company_values:
        company_box.current(0)

    tb.Label(filters_frame, text="Fecha inicio").grid(row=0, column=3, sticky="w")
    start_var = DateEntry(filters_frame, dateformat="%Y-%m-%d")
    start_var.grid(row=1, column=3, sticky="we", padx=(0, 10))
    start_var.entry.delete(0, tk.END)

    tb.Label(filters_frame, text="Fecha fin").grid(row=0, column=4, sticky="w")
    end_var = DateEntry(filters_frame, dateformat="%Y-%m-%d")
    end_var.grid(row=1, column=4, sticky="we", padx=(0, 10))
    end_var.entry.delete(0, tk.END)

    tb.Label(filters_frame, text="Buscar").grid(row=0, column=5, sticky="w")
    search_var = tk.StringVar(value="")
    search_entry = tb.Entry(filters_frame, textvariable=search_var)
    search_entry.grid(row=1, column=5, sticky="we")

    filters_frame.grid_columnconfigure(5, weight=1)

    best_filter_var = tk.StringVar(value="Todas")
    dde_filter_var = tk.StringVar(value="Todas")

    table_frame = tb.Frame(parent, padding=(12, 0))
    table_frame.pack(fill=BOTH, expand=YES)

    columns_config: Dict[str, Dict[str, object]] = {
        "ticket": {"text": "Ticket", "width": 110, "anchor": "center"},
        "titulo": {"text": "Titulo", "width": 360, "stretch": True},
        "tipo": {"text": "Tipo incidente", "width": 180},
        "status": {"text": "Status", "width": 140},
        "empresa": {"text": "Empresa", "width": 200},
        "actualizado": {"text": "Actualizado", "width": 160, "anchor": "center"},
        "mejor": {"text": "Mejor respuesta", "width": 140, "anchor": "center"},
        "dde": {"text": "DDE generada", "width": 140, "anchor": "center"},
    }
    columns = tuple(columns_config.keys())
    tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=14)
    tree["displaycolumns"] = columns

    active_sort: Dict[str, Optional[str]] = {"column": None, "direction": None}

    def _coerce_sort_value(value: str, column: str) -> object:
        """Return a comparable value for sorting operations."""

        if column == "actualizado" and value:
            try:
                return datetime.strptime(value, "%Y-%m-%d %H:%M")
            except ValueError:
                return value
        if column == "ticket":
            try:
                return int(value)
            except (TypeError, ValueError):
                return value
        normalized = str(value).strip().lower()
        if normalized in {"si", "no"}:
            return 1 if normalized == "si" else 0
        try:
            return float(value)
        except (TypeError, ValueError):
            return normalized

    def _sort_tree(column: str, force_direction: Optional[str] = None) -> None:
        """Sort the tree rows using the provided column."""

        direction = force_direction
        if direction is None:
            if active_sort["column"] == column and active_sort["direction"] == "asc":
                direction = "desc"
            else:
                direction = "asc"
        active_sort["column"] = column
        active_sort["direction"] = direction
        items = [(tree.set(item_id, column), item_id) for item_id in tree.get_children("")]
        reverse = direction == "desc"
        try:
            items.sort(key=lambda item: _coerce_sort_value(item[0], column), reverse=reverse)
        except TypeError:
            items.sort(key=lambda item: str(item[0]).lower(), reverse=reverse)
        for index, (_, item_id) in enumerate(items):
            tree.move(item_id, "", index)

    for column_name, column_config in columns_config.items():
        tree.heading(column_name, text=column_config["text"], command=lambda col=column_name: _sort_tree(col))
        tree.column(
            column_name,
            width=int(column_config.get("width", 120)),
            anchor=column_config.get("anchor", tk.W),
            stretch=bool(column_config.get("stretch", False)),
        )
    tree.tag_configure("best", background="#e6f4ea")
    tree.tag_configure("dde", background="#fff7e6")

    scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=scrollbar.set)
    tree.pack(side=LEFT, fill=BOTH, expand=YES)
    scrollbar.pack(side=RIGHT, fill=Y)
    bind_mousewheel(tree, tree.yview)

    actions_frame = tb.Frame(parent, padding=12)
    actions_frame.pack(fill=X)
    status_label = tb.Label(actions_frame, text="Selecciona una tarjeta para comenzar.", bootstyle=SECONDARY)
    status_label.pack(side=LEFT)

    column_vars: Dict[str, tk.BooleanVar] = {}

    def _toggle_column(column: str) -> None:
        """Hide or show the requested column ensuring at least one stays visible."""

        visible_columns = [name for name in columns if column_vars[name].get()]
        if not visible_columns:
            column_vars[column].set(True)
            messagebox.showinfo("Columnas", "Debe permanecer al menos una columna visible.")
            return
        tree["displaycolumns"] = visible_columns

    column_button = ttk.Menubutton(actions_frame, text="Columnas")
    column_menu = tk.Menu(column_button, tearoff=False)
    column_button["menu"] = column_menu
    for column_name, column_config in columns_config.items():
        var = tk.BooleanVar(value=True)
        column_vars[column_name] = var
        column_menu.add_checkbutton(
            label=column_config["text"],
            variable=var,
            command=lambda col=column_name: _toggle_column(col),
        )
    history_button = tb.Button(
        actions_frame,
        text="Historial",
        bootstyle=INFO,
        state=tk.DISABLED,
        command=lambda: _show_history(root, controller, selected_card[0].cardId)
        if selected_card
        else None,
    )
    generate_button = tb.Button(
        actions_frame,
        text="Generar DDE/HU",
        bootstyle=PRIMARY,
        state=tk.DISABLED,
        command=lambda: _open_capture_form(root, controller, selected_card[0]) if selected_card else None,
    )
    generate_button.pack(side=RIGHT)
    history_button.pack(side=RIGHT, padx=(0, 6))
    column_button.pack(side=RIGHT, padx=(0, 6))

    selected_card: List[CardDTO] = []
    current_cards: List[CardDTO] = []

    def _parse_date(entry: DateEntry) -> Optional[datetime]:
        """Return the selected date or ``None`` when the field is empty."""

        value = entry.entry.get().strip()
        if not value:
            return None
        try:
            return datetime.strptime(value + " 00:00", "%Y-%m-%d %H:%M")
        except ValueError:
            return None

    def _refresh() -> None:
        """Load the cards from the controller applying filters."""

        selected_type = tipo_var.get().strip()
        selected_status = status_var.get().strip()
        selected_company = company_var.get().strip()
        filters: Dict[str, object] = {
            "fechaInicio": _parse_date(start_var),
            "fechaFin": _parse_date(end_var),
            "busqueda": search_var.get().strip() or None,
            "estadoMejor": best_filter_var.get(),
            "estadoDde": dde_filter_var.get(),
        }
        if selected_type:
            filters["tipoId"] = incident_type_map.get(selected_type)
        if selected_status:
            filters["status"] = selected_status
        if selected_company:
            filters["empresaId"] = company_map.get(selected_company)
        try:
            cards = controller.list_cards(filters)
        except RuntimeError as exc:
            messagebox.showerror("Error", str(exc))
            return

        current_cards[:] = cards
        selected_card.clear()
        generate_button.configure(state=tk.DISABLED)
        history_button.configure(state=tk.DISABLED)
        tree.delete(*tree.get_children(""))
        for card in cards:
            tags_list: List[str] = []
            if card.hasBestSelection:
                tags_list.append("best")
            if card.hasDdeGenerated:
                tags_list.append("dde")
            formatted_date = _format_datetime(card.updatedAt or card.createdAt)
            tree.insert(
                "",
                "end",
                iid=str(card.cardId),
                values=(
                    card.ticketId or str(card.cardId),
                    card.title,
                    card.incidentTypeName or "",
                    card.status,
                    card.companyName or "",
                    formatted_date,
                    "Si" if card.hasBestSelection else "No",
                    "Si" if card.hasDdeGenerated else "No",
                ),
                tags=tuple(tags_list),
            )
        if active_sort["column"] and active_sort["direction"]:
            _sort_tree(active_sort["column"], force_direction=active_sort["direction"])
        status_label.configure(text=f"{len(cards)} tarjeta(s) encontradas.")

    def _on_select(event: tk.Event) -> None:
        """Enable the action buttons when a card is selected."""

        selection = tree.selection()
        if not selection:
            selected_card.clear()
            generate_button.configure(state=tk.DISABLED)
            return
        card_id = int(selection[0])
        for card in current_cards:
            if card.cardId == card_id:
                selected_card[:] = [card]
                break
        state = tk.NORMAL if selected_card else tk.DISABLED
        generate_button.configure(state=state)
        history_button.configure(state=state)

    tree.bind("<<TreeviewSelect>>", _on_select)

    debounce_id = None

    def _schedule_refresh(*_args: object) -> None:
        """Apply debounce to the search entry."""

        nonlocal debounce_id
        if debounce_id is not None:
            parent.after_cancel(debounce_id)
        debounce_id = parent.after(300, _refresh)

    tb.Label(filters_frame, text="Mejor respuesta").grid(row=2, column=0, sticky="w", pady=(8, 0))
    best_filter_box = ttk.Combobox(
        filters_frame,
        values=("Todas", "Con mejor respuesta", "Sin mejor respuesta"),
        textvariable=best_filter_var,
        state="readonly",
        width=22,
    )
    best_filter_box.grid(row=3, column=0, sticky="we", padx=(0, 10))
    best_filter_box.current(0)

    tb.Label(filters_frame, text="DDE generada").grid(row=2, column=1, sticky="w", pady=(8, 0))
    dde_filter_box = ttk.Combobox(
        filters_frame,
        values=("Todas", "Con DDE generada", "Sin DDE generada"),
        textvariable=dde_filter_var,
        state="readonly",
        width=22,
    )
    dde_filter_box.grid(row=3, column=1, sticky="we", padx=(0, 10))
    dde_filter_box.current(0)

    for widget in (tipo_box, status_box, company_box):
        widget.bind("<<ComboboxSelected>>", lambda *_: _refresh(), add="+")
    start_var.bind("<<DateEntrySelected>>", lambda *_: _refresh(), add="+")
    end_var.bind("<<DateEntrySelected>>", lambda *_: _refresh(), add="+")
    search_var.trace_add("write", lambda *_: _schedule_refresh())
    best_filter_box.bind("<<ComboboxSelected>>", lambda *_: _refresh(), add="+")
    dde_filter_box.bind("<<ComboboxSelected>>", lambda *_: _refresh(), add="+")

    _refresh()

