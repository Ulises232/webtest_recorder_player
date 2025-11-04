"""Business helpers to export AI documents into structured directories."""

from __future__ import annotations

import json
import re
from enum import Enum
from html import escape
from importlib import resources
from pathlib import Path
from typing import Dict, Optional

from app.config.storage_paths import getDdeExportBaseDirectory
from app.dtos.card_ai_dto import CardAIOutputDTO, CardDTO

HTML_TEMPLATE_PACKAGE = "app.templates"
HTML_TEMPLATE_NAME = "card_generation.html"


class CardAIExportServiceError(RuntimeError):
    """Raised when an exported document cannot be generated."""


class CardAIExportFormat(str, Enum):
    """Enumerate the supported export formats."""

    JSON = "json"
    MARKDOWN = "markdown"
    DOCX = "docx"
    HTML = "html"

    @property
    def extension(self) -> str:
        """Return the file extension associated with the export format."""

        return {
            CardAIExportFormat.JSON: ".json",
            CardAIExportFormat.MARKDOWN: ".md",
            CardAIExportFormat.DOCX: ".docx",
            CardAIExportFormat.HTML: ".html",
        }[self]


class CardAIExportService:
    """Persist AI outputs under the folder structure required by the business."""

    def __init__(
        self,
        base_directory: Optional[Path] = None,
        templates_package: str = HTML_TEMPLATE_PACKAGE,
        html_template_name: str = HTML_TEMPLATE_NAME,
    ) -> None:
        """Store the base directory and template parameters."""

        self._base_directory = base_directory
        self._templates_package = templates_package
        self._html_template_name = html_template_name

    def export_output(
        self,
        card: CardDTO,
        output: CardAIOutputDTO,
        export_format: CardAIExportFormat,
    ) -> Path:
        """Export the output content to disk and return the generated path."""

        content = output.content
        if not isinstance(content, dict):
            raise CardAIExportServiceError("El contenido de la tarjeta debe ser un objeto JSON.")

        target_directory = self._build_target_directory(card)
        target_directory.mkdir(parents=True, exist_ok=True)

        filename = self._build_filename(card, export_format)
        destination = target_directory / filename

        try:
            if export_format is CardAIExportFormat.JSON:
                self._write_json(destination, content)
            elif export_format is CardAIExportFormat.MARKDOWN:
                self._write_markdown(destination, content)
            elif export_format is CardAIExportFormat.DOCX:
                self._write_docx(destination, content)
            elif export_format is CardAIExportFormat.HTML:
                self._write_html(destination, content)
            else:  # pragma: no cover - Enum exhaustive guard
                raise CardAIExportServiceError("Formato de exportación desconocido.")
        except OSError as exc:
            raise CardAIExportServiceError(f"No fue posible guardar el archivo destino: {exc}") from exc

        return destination

    def _resolve_base_directory(self) -> Path:
        """Return the base directory where the documents should be stored."""

        if self._base_directory is not None:
            return self._base_directory
        return getDdeExportBaseDirectory(create=True)

    def _build_target_directory(self, card: CardDTO) -> Path:
        """Return the directory path that corresponds to the card metadata."""

        base_directory = self._resolve_base_directory()
        company_segment = self._sanitize_segment(card.companyName or "Sin_empresa")
        path = base_directory / company_segment

        sprint_segment = ""
        if card.sprintName and card.sprintName.strip():
            sprint_segment = self._sanitize_segment(card.sprintName)
        if sprint_segment:
            path = path / sprint_segment

        return path

    @staticmethod
    def _sanitize_segment(segment: str) -> str:
        """Transform arbitrary text into a safe path segment."""

        cleaned = segment.strip()
        cleaned = cleaned.replace(" ", "_")
        cleaned = re.sub(r"[<>:\"/\\\\|?*]", "_", cleaned)
        cleaned = re.sub(r"_+", "_", cleaned)
        cleaned = cleaned.strip("._")
        return cleaned or "valor"

    def _build_filename(self, card: CardDTO, export_format: CardAIExportFormat) -> str:
        """Return the filename (without directory) for the export."""

        ticket = (card.ticketId or "").strip()
        if not ticket:
            ticket = f"card_{card.cardId}"
        ticket = ticket.replace(" ", "_")
        ticket = re.sub(r"[<>:\"/\\\\|?*]", "-", ticket)
        ticket = re.sub(r"-+", "-", ticket).strip("-")
        if not ticket:
            ticket = f"card_{card.cardId}"
        return f"{ticket}{export_format.extension}"

    @staticmethod
    def _write_json(destination: Path, content: Dict[str, object]) -> None:
        """Persist the JSON payload with indentation."""

        with destination.open("w", encoding="utf-8") as handle:
            json.dump(content, handle, ensure_ascii=False, indent=2)

    @staticmethod
    def _render_markdown(content: Dict[str, object]) -> str:
        """Convert the JSON document into Markdown."""

        lines = ["# Documento generado"]
        for key, value in content.items():
            lines.append(f"\n## {key}")
            if isinstance(value, list):
                for item in value:
                    lines.append(f"- {item}")
            else:
                lines.append(str(value))
        return "\n".join(lines)

    def _write_markdown(self, destination: Path, content: Dict[str, object]) -> None:
        """Write the Markdown representation to disk."""

        document = self._render_markdown(content)
        with destination.open("w", encoding="utf-8") as handle:
            handle.write(document)

    def _write_docx(self, destination: Path, content: Dict[str, object]) -> None:
        """Serialize the content to DOCX format."""

        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependency missing
            raise CardAIExportServiceError(
                "La dependencia 'python-docx' es requerida para exportar en formato DOCX."
            ) from exc

        document = Document()
        document.add_heading("Documento DDE/HU", level=1)
        for key, value in content.items():
            document.add_heading(str(key), level=2)
            if isinstance(value, list):
                for item in value:
                    document.add_paragraph(str(item), style="List Bullet")
            else:
                document.add_paragraph(str(value))
        document.save(destination)

    def _write_html(self, destination: Path, content: Dict[str, object]) -> None:
        """Render the HTML template filled with the JSON content."""

        template = self._load_html_template()
        rendered = self._fill_html_template(template, content)
        with destination.open("w", encoding="utf-8") as handle:
            handle.write(rendered)

    def _load_html_template(self) -> str:
        """Return the HTML template content from the configured resource package."""

        try:
            template_resource = resources.files(self._templates_package).joinpath(
                self._html_template_name
            )
        except ModuleNotFoundError as exc:  # pragma: no cover - invalid configuration
            raise CardAIExportServiceError("No se encontró el paquete de plantillas HTML.") from exc

        try:
            with template_resource.open("r", encoding="utf-8") as handle:
                return handle.read()
        except FileNotFoundError as exc:
            raise CardAIExportServiceError("No se encontró la plantilla HTML configurada.") from exc
        except OSError as exc:  # pragma: no cover - filesystem errors
            raise CardAIExportServiceError(
                "Ocurrió un error al leer la plantilla HTML configurada."
            ) from exc

    @staticmethod
    def _fill_html_template(template: str, content: Dict[str, object]) -> str:
        """Inject the JSON values into the HTML placeholders."""

        placeholders = {
            "titulo": CardAIExportService._format_html_value(content.get("titulo")),
            "tipo": CardAIExportService._format_html_value(content.get("tipo")),
            "fecha": CardAIExportService._format_html_value(content.get("fecha")),
            "hora_inicio": CardAIExportService._format_html_value(content.get("hora_inicio")),
            "hora_fin": CardAIExportService._format_html_value(content.get("hora_fin")),
            "lugar": CardAIExportService._format_html_value(content.get("lugar")),
            "descripcion": CardAIExportService._format_html_value(content.get("descripcion")),
            "que_necesitas": CardAIExportService._format_html_value(content.get("que_necesitas")),
            "para_que_lo_necesitas": CardAIExportService._format_html_value(
                content.get("para_que_lo_necesitas")
            ),
            "como_lo_necesitas": CardAIExportService._format_html_value(
                content.get("como_lo_necesitas")
            ),
            "requerimientos_funcionales": CardAIExportService._format_html_value(
                content.get("requerimientos_funcionales")
            ),
            "requerimientos_especiales": CardAIExportService._format_html_value(
                content.get("requerimientos_especiales")
            ),
            "criterios_aceptacion": CardAIExportService._format_html_value(
                content.get("criterios_aceptacion")
            ),
        }

        rendered = template
        for key, value in placeholders.items():
            rendered = rendered.replace(f"{{{{{key}}}}}", value)

        return re.sub(r"\{\{[a-zA-Z0-9_]+\}\}", "&nbsp;", rendered)

    @staticmethod
    def _format_html_value(value: object) -> str:
        """Return a string ready to be interpolated within the HTML template."""

        if isinstance(value, list):
            items = [escape(str(item)) for item in value if str(item).strip()]
            if not items:
                return "&nbsp;"
            return "<ul>" + "".join(f"<li>{item}</li>" for item in items) + "</ul>"

        if value is None:
            return "&nbsp;"

        text = escape(str(value).strip())
        if not text:
            return "&nbsp;"
        return text.replace("\n", "<br />")
